"""Generate CNN_ViT_CIFAR10_Report.docx — research-style report for the 4-experiment study.

Updated with real Colab training results (March 2026).

Usage:
    python report/generate_report.py
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PLOT_DIR = os.path.join(BASE_DIR, "plots")
OUT_PATH = os.path.join(BASE_DIR, "report", "CNN_ViT_CIFAR10_Report.docx")

TITLE_CLR = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
HDR_BG = "1F3864"
ALT_BG = "EBF0F8"


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = TITLE_CLR if level == 1 else ACCENT
    for run in p.runs:
        run.font.color.rgb = color
    return p


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def make_table(doc, headers, rows, caption=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_bg(cell, HDR_BG)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.bold = True

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            if i % 2 == 1:
                _set_cell_bg(cell, ALT_BG)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.italic = True

    return table


def add_image(doc, filename, width=5.5, caption=None):
    path = os.path.join(PLOT_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True


def build_report():
    doc = Document()

    # ── Title page ──────────────────────────────────────────────────────
    title = doc.add_heading("CNN and Vision Transformers on CIFAR-10", level=0)
    for run in title.runs:
        run.font.color.rgb = TITLE_CLR
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "A Systematic Architecture Comparison:\n"
        "Baseline CNN \u2192 Residual CNN \u2192 Vision Transformer \u2192 Hybrid CNN-ViT"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(13)

    author = doc.add_paragraph("Tanvir")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author.runs:
        run.font.size = Pt(12)
        run.italic = True

    doc.add_paragraph()

    # ═════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "1. Introduction")
    body(doc,
         "This report presents a systematic experimental study comparing four deep learning "
         "architectures on the CIFAR-10 image classification benchmark. All four models were "
         "trained from scratch on Google Colab using an NVIDIA GPU with mixed-precision (AMP) "
         "training. The study progressively builds from a simple VGG-style CNN baseline to a "
         "Hybrid CNN-Vision Transformer model, evaluating the impact of architectural innovations "
         "such as residual connections, self-attention mechanisms, and CNN-Transformer integration.")

    body(doc,
         "The four experiments are: (1) a Baseline VGG-style CNN achieving 92.15% accuracy, "
         "(2) a Residual CNN with skip connections achieving 92.76%, (3) a pure Vision Transformer "
         "(ViT) achieving 69.55%, and (4) a Hybrid CNN + Vision Transformer achieving 90.95%. "
         "All experiments use identical training conditions (AdamW optimizer, OneCycleLR scheduler, "
         "EMA, label smoothing) to ensure that accuracy differences reflect only architectural changes.")

    heading(doc, "1.1 Thought Process and Approach", level=2)
    body(doc,
         "The idea behind this study was to explore how different architectural paradigms handle "
         "the same image classification task. CNNs have been the dominant architecture for computer "
         "vision due to their built-in inductive biases (translation equivariance, locality). Vision "
         "Transformers (ViT) recently challenged this by showing that pure attention mechanisms can "
         "match or exceed CNN performance on large-scale datasets. The key question motivating this "
         "study: Can we combine the strengths of both architectures for better performance on a "
         "small-scale dataset like CIFAR-10?")

    body(doc,
         "The experimental progression was designed deliberately: start with the simplest CNN "
         "(Experiment 1) to establish a baseline, then add residual connections (Experiment 2) to "
         "see how skip connections improve gradient flow, then try a pure Transformer (Experiment 3) "
         "to understand its limitations on small images, and finally combine CNN + Transformer "
         "(Experiment 4) to test the hybrid hypothesis.")

    # ═════════════════════════════════════════════════════════════════════
    # 2. DATASET DESCRIPTION
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "2. Dataset Description")
    body(doc,
         "CIFAR-10 is a widely-used image classification benchmark consisting of 60,000 colour "
         "images (32\u00d732 pixels, RGB) across 10 mutually exclusive classes: airplane, automobile, "
         "bird, cat, deer, dog, frog, horse, ship, and truck.")

    make_table(doc,
               headers=["Property", "Value"],
               rows=[
                   ["Total images", "60,000"],
                   ["Training set", "50,000"],
                   ["Test set", "10,000"],
                   ["Image size", "32 x 32 x 3 (RGB)"],
                   ["Number of classes", "10"],
                   ["Images per class", "6,000 (balanced)"],
               ],
               caption="Table 1. CIFAR-10 dataset summary.")

    add_image(doc, "sample_images.png", 5.5, "Figure 1. CIFAR-10 sample images from all 10 classes.")

    heading(doc, "2.1 Data Augmentation", level=2)
    body(doc,
         "Data augmentation was critical for preventing overfitting on the relatively small "
         "training set. The following pipeline was applied to training data only:")
    make_table(doc,
               headers=["Transform", "Parameters", "Purpose"],
               rows=[
                   ["RandomCrop", "32, padding=4", "Translation invariance via random spatial shifts"],
                   ["RandomHorizontalFlip", "p=0.5", "Mirror invariance for symmetric objects"],
                   ["RandAugment", "num_ops=2, magnitude=9", "Automated diverse augmentations"],
                   ["Normalize", "mean/std per-channel", "Standardize input distribution"],
               ],
               caption="Table 2. Training augmentation pipeline.")

    body(doc,
         "Reasoning: RandAugment was chosen over manual augmentation because it automatically "
         "applies a diverse set of transformations (rotation, color jitter, sharpness, etc.) "
         "without requiring manual tuning of each augmentation hyperparameter. The magnitude "
         "of 9 (out of 30) provides moderate augmentation strength suitable for CIFAR-10's "
         "small image size.")

    # ═════════════════════════════════════════════════════════════════════
    # 3. METHODOLOGY
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "3. Methodology")
    body(doc,
         "A shared training pipeline ensures fair comparison across all experiments. "
         "Keeping the training recipe constant isolates the effect of architecture alone.")
    make_table(doc,
               headers=["Component", "Setting", "Rationale"],
               rows=[
                   ["Optimizer", "AdamW (lr=3e-4, wd=0.05)", "Decoupled weight decay; stable for both CNNs and Transformers"],
                   ["Scheduler", "OneCycleLR", "Linear warmup prevents early divergence; cosine annealing finds flatter minima"],
                   ["EMA", "decay=0.999", "Smoothes weight updates, reduces validation noise"],
                   ["Gradient clipping", "max_norm=1.0", "Prevents gradient explosions in Transformer attention layers"],
                   ["Label smoothing", "0.1", "Reduces overconfident predictions, improves calibration"],
                   ["Batch size", "128", "Balances GPU utilization and gradient noise"],
                   ["Loss function", "CrossEntropyLoss", "Standard multi-class classification loss"],
                   ["Mixed precision", "AMP (CUDA only)", "2x training speed with FP16 forward pass"],
               ],
               caption="Table 3. Shared training configuration.")

    body(doc,
         "Reasoning: AdamW was selected over SGD because it provides more stable convergence "
         "for Transformer architectures (Experiment 3 and 4) while still working well for CNNs. "
         "The learning rate of 3e-4 follows the common default for AdamW on vision tasks. Weight "
         "decay of 0.05 provides regularization proportional to parameter magnitude. OneCycleLR "
         "was chosen because it automatically handles warmup and decay, which is especially "
         "important for Transformers that are sensitive to learning rate warmup.")

    # ═════════════════════════════════════════════════════════════════════
    # 4. EXPERIMENT 1 — BASELINE CNN
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "4. Experiment 1 \u2014 Baseline CNN")

    heading(doc, "4.1 Thought Process", level=2)
    body(doc,
         "The first experiment establishes a baseline using the most straightforward CNN design: "
         "a VGG-style architecture with stacked 3\u00d73 convolutions. The reasoning was to start with "
         "the simplest possible competitive architecture, so that improvements in later experiments "
         "can be clearly attributed to specific architectural changes. VGG-style networks are easy "
         "to understand and implement, making them an ideal reference point.")

    body(doc,
         "Key design decisions: (1) Three convolutional blocks with increasing channel depth "
         "(64 \u2192 128 \u2192 256) to progressively extract more abstract features. (2) MaxPooling after "
         "each block to reduce spatial dimensions. (3) A flat fully-connected (FC) classification "
         "head, which is the traditional approach but has a known weakness of containing many parameters.")

    heading(doc, "4.2 Architecture", level=2)
    make_table(doc,
               headers=["Layer", "Output Shape", "Description"],
               rows=[
                   ["Input", "3 x 32 x 32", "RGB image"],
                   ["Block 1 (2 conv)", "64 x 16 x 16", "Conv(3\u219264), BN, ReLU x2 + MaxPool"],
                   ["Block 2 (2 conv)", "128 x 8 x 8", "Conv(64\u2192128), BN, ReLU x2 + MaxPool"],
                   ["Block 3 (1 conv)", "256 x 4 x 4", "Conv(128\u2192256), BN, ReLU + MaxPool"],
                   ["Flatten", "4,096", "256 * 4 * 4"],
                   ["FC + Dropout(0.5)", "512", "Linear(4096\u2192512), ReLU, Dropout"],
                   ["Classifier", "10", "Linear(512\u219210)"],
               ],
               caption="Table 4. Baseline CNN architecture (2,658,762 trainable parameters).")

    heading(doc, "4.3 Training Results", level=2)
    make_table(doc,
               headers=["Metric", "Value"],
               rows=[
                   ["Trainable parameters", "2,658,762"],
                   ["Training epochs", "100"],
                   ["Best val accuracy", "92.15%"],
                   ["Final train accuracy", "93.2%"],
                   ["Training time", "~68 minutes (4,087s)"],
               ],
               caption="Table 5. Experiment 1 results (real Colab training).")

    add_image(doc, "exp1_curves.png", 5.0, "Figure 2. Experiment 1: Baseline CNN training curves (loss, accuracy, LR).")

    body(doc,
         "Analysis: The baseline CNN achieved a surprisingly strong 92.15% validation accuracy, "
         "largely thanks to the modern training recipe (AdamW + OneCycleLR + EMA + label smoothing + "
         "RandAugment). The learning rate schedule shows clear OneCycleLR warmup (epochs 1\u201330) "
         "followed by cosine annealing. The generalization gap (train 93.2% vs val 92.1%) is small, "
         "indicating good regularization from Dropout(0.5), label smoothing, and augmentation. However, "
         "the flat FC head contains over 2M of the 2.7M total parameters, making it parameter-inefficient.")

    # ═════════════════════════════════════════════════════════════════════
    # 5. EXPERIMENT 2 — IMPROVED CNN
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "5. Experiment 2 \u2014 Improved CNN (Residual CNN)")

    heading(doc, "5.1 Thought Process", level=2)
    body(doc,
         "After seeing the baseline CNN results, the next step was to address its known limitations: "
         "(1) No skip connections, meaning gradients must flow through every layer sequentially, which "
         "can cause vanishing gradients in deeper networks. (2) The flat FC head is over-parameterized. "
         "(3) Only three convolutional stages limit representational capacity.")

    body(doc,
         "The idea was inspired by ResNet (He et al., 2015): add residual connections that let the "
         "network learn residual functions R(x) = F(x) - x instead of the full mapping F(x). This "
         "is mathematically easier to optimize because the identity mapping provides a baseline that "
         "the network can refine. We also replaced the flat FC head with Global Average Pooling (GAP), "
         "which directly averages spatial features into a single vector, eliminating millions of "
         "parameters and acting as a structural regularizer.")

    heading(doc, "5.2 Architecture", level=2)
    make_table(doc,
               headers=["Layer", "Output Shape", "Description"],
               rows=[
                   ["Stem", "64 x 32 x 32", "Conv(3\u219264), BN, ReLU"],
                   ["ResBlock 1", "128 x 16 x 16", "Conv\u2192BN\u2192ReLU\u2192Conv\u2192BN + skip (stride=2)"],
                   ["ResBlock 2", "256 x 8 x 8", "Conv\u2192BN\u2192ReLU\u2192Conv\u2192BN + skip (stride=2)"],
                   ["ResBlock 3", "256 x 8 x 8", "Conv\u2192BN\u2192ReLU\u2192Conv\u2192BN + skip (stride=1)"],
                   ["ResBlock 4", "512 x 4 x 4", "Conv\u2192BN\u2192ReLU\u2192Conv\u2192BN + skip (stride=2)"],
                   ["GAP + Dropout(0.3)", "512", "AdaptiveAvgPool2d(1) + Flatten + Dropout"],
                   ["Classifier", "10", "Linear(512\u219210)"],
               ],
               caption="Table 6. Residual CNN architecture (6,009,930 trainable parameters).")

    heading(doc, "5.3 Training Results", level=2)
    make_table(doc,
               headers=["Metric", "Baseline CNN", "Residual CNN", "Change"],
               rows=[
                   ["Val accuracy", "92.15%", "92.76%", "+0.61%"],
                   ["Parameters", "2,658,762", "6,009,930", "+126% more"],
                   ["Epochs trained", "100", "60", "40% fewer epochs"],
                   ["Training time", "4,087s", "2,520s", "38% faster"],
                   ["Final train acc", "93.2%", "96.6%", "+3.4%"],
                   ["Pooling method", "Flat FC", "GAP", "More efficient"],
               ],
               caption="Table 7. Experiment 2 vs Experiment 1 comparison.")

    add_image(doc, "exp2_curves.png", 5.0, "Figure 3. Experiment 2: Residual CNN training curves.")

    body(doc,
         "Analysis: The Residual CNN achieved the highest accuracy of all four experiments at 92.76%, "
         "reaching this in only 60 epochs (vs 100 for the baseline). The faster convergence is directly "
         "attributable to residual connections enabling better gradient flow. The model converges to "
         ">90% val acc by epoch 30, while the baseline needed ~60 epochs. However, the larger "
         "generalization gap (train 96.6% vs val 92.7% = 3.9%) suggests the deeper network has more "
         "capacity than needed for CIFAR-10, with some overfitting in later epochs.")

    # ═════════════════════════════════════════════════════════════════════
    # 6. EXPERIMENT 3 — VISION TRANSFORMER
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "6. Experiment 3 \u2014 Vision Transformer (ViT)")

    heading(doc, "6.1 Thought Process", level=2)
    body(doc,
         "With two strong CNN baselines established, the next question was: Can a pure Transformer "
         "architecture, which has revolutionized NLP and shown strong results on large-scale image "
         "datasets (ImageNet), work well on the small-scale CIFAR-10? Vision Transformers (Dosovitskiy "
         "et al., 2020) process images by dividing them into patches, linearly embedding each patch, "
         "and applying standard Transformer encoder blocks with self-attention.")

    body(doc,
         "The hypothesis was that ViTs might struggle on CIFAR-10 because: (1) The images are only "
         "32\u00d732 pixels, severely limiting the number of meaningful patches. (2) With only 50,000 "
         "training images, there may not be enough data for the Transformer to learn low-level visual "
         "features that CNNs get for free from their convolutional inductive bias. (3) Self-attention "
         "has O(n\u00b2) complexity, which is expensive even with 64 patches.")

    heading(doc, "6.2 Self-Attention Mechanism", level=2)
    body(doc,
         "Self-attention computes pairwise relationships between all patches simultaneously. "
         "For each patch token, the mechanism produces Query (Q), Key (K), and Value (V) vectors. "
         "Attention weights are computed as softmax(QK^T / sqrt(d_head)), then used to weight "
         "the Values. Multiple heads attend to different representational subspaces, enabling "
         "the model to capture both spatial proximity and semantic similarity patterns.")

    heading(doc, "6.3 Architecture", level=2)
    make_table(doc,
               headers=["Component", "Configuration"],
               rows=[
                   ["Patch size", "4 x 4 (64 tokens from 32x32)"],
                   ["Embedding dimension", "192"],
                   ["Transformer layers", "6"],
                   ["Attention heads", "3 (head_dim = 64)"],
                   ["MLP ratio", "4.0 (hidden_dim = 768)"],
                   ["Stochastic depth", "linearly 0 \u2192 0.1"],
                   ["Dropout", "0.1"],
                   ["Positional embedding", "Learnable (not sinusoidal)"],
                   ["Classification", "Mean pooling over tokens + Linear"],
               ],
               caption="Table 8. ViT configuration (2,693,194 trainable parameters).")

    heading(doc, "6.4 Training Results", level=2)
    make_table(doc,
               headers=["Metric", "Value"],
               rows=[
                   ["Trainable parameters", "2,693,194"],
                   ["Training epochs", "40"],
                   ["Best val accuracy", "69.55%"],
                   ["Final train accuracy", "64.0%"],
                   ["Training time", "~31 minutes (1,845s)"],
               ],
               caption="Table 9. Experiment 3 results (real Colab training).")

    add_image(doc, "exp3_curves.png", 5.0, "Figure 4. Experiment 3: Vision Transformer training curves.")

    body(doc,
         "Analysis: The pure ViT achieved only 69.55% validation accuracy \u2014 dramatically "
         "lower than both CNN models (~92%). This confirms the well-known finding that Vision "
         "Transformers struggle on small datasets without pre-training. Key observations:")

    bullet(doc,
           "Slow convergence: After 40 epochs, the model was still improving but had not "
           "plateaued, suggesting it needs significantly more training time than CNNs.")
    bullet(doc,
           "Low training accuracy (64.0%): Even on training data, the model underfits, "
           "confirming that the Transformer cannot learn low-level features efficiently from "
           "50K small images alone.")
    bullet(doc,
           "The ViT must learn spatial locality from scratch, while CNNs have this built in "
           "via their convolutional structure. On 32\u00d732 images with only 64 patches, there is "
           "insufficient spatial resolution for attention to discover useful patterns quickly.")

    heading(doc, "6.5 Why ViT Failed on CIFAR-10", level=2)
    body(doc,
         "The poor performance is not a flaw of the Transformer architecture itself, but a "
         "consequence of the data scale. The original ViT paper showed that ViTs need large-scale "
         "pre-training (ImageNet-21k or JFT-300M) to outperform CNNs. On CIFAR-10 alone, the "
         "Transformer spends most of its capacity learning basic edge detection and spatial "
         "relationships that CNNs encode structurally. With 300 epochs and stronger augmentation, "
         "the ViT could potentially reach ~80%+, but it would still lag behind CNNs on this scale.")

    # ═════════════════════════════════════════════════════════════════════
    # 7. EXPERIMENT 4 — HYBRID CNN-ViT
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "7. Experiment 4 \u2014 Hybrid CNN + Vision Transformer")

    heading(doc, "7.1 Thought Process", level=2)
    body(doc,
         "The ViT's poor performance in Experiment 3 motivated a hybrid approach: what if we use a "
         "CNN to extract local features first, then let a Transformer capture global relationships? "
         "This was inspired by several observations:")

    bullet(doc,
           "CNNs excel at low-level feature extraction (edges, textures, shapes) due to their "
           "convolutional inductive bias, but are limited by their local receptive field.")
    bullet(doc,
           "Transformers excel at capturing global relationships via self-attention, but waste "
           "capacity learning basic visual features from raw pixels.")
    bullet(doc,
           "A CNN 'stem' can pre-process raw pixels into meaningful feature maps, giving the "
           "Transformer a head start and injecting the local inductive bias it lacks.")

    body(doc,
         "The hypothesis was that combining CNN local feature extraction + Transformer global "
         "attention would outperform either approach alone. The CNN stem would handle what CNNs "
         "do best (local patterns), while the Transformer would handle what it does best (global "
         "context and long-range dependencies).")

    heading(doc, "7.2 CNN Stem Design", level=2)
    body(doc,
         "The CNN stem consists of two convolutional layers with BatchNorm and GELU activation. "
         "It preserves spatial resolution (32\u00d732) while expanding the channel dimension from 3 "
         "to 128. GELU was chosen over ReLU because it provides smoother gradients, which works "
         "better with the downstream Transformer layers. The stem provides the Transformer with "
         "locally-aware feature maps instead of raw pixel patches.")

    heading(doc, "7.3 Transformer Encoder", level=2)
    body(doc,
         "After the CNN stem extracts local features, the feature maps are divided into 4\u00d74 "
         "patches (64 tokens) and linearly projected to 256 dimensions. The Transformer encoder "
         "consists of 6 blocks, each with 8-head self-attention and a 4x MLP. PreNorm architecture "
         "(LayerNorm before attention/MLP) and StochasticDepth regularization ensure stable training.")

    heading(doc, "7.4 Architecture", level=2)
    make_table(doc,
               headers=["Stage", "Component", "Output Shape"],
               rows=[
                   ["Input", "RGB image", "3 x 32 x 32"],
                   ["CNN Stem", "Conv(3\u219264)+BN+GELU, Conv(64\u2192128)+BN+GELU", "128 x 32 x 32"],
                   ["Patch Embed", "Conv(128\u2192256, k=4, s=4) + pos_embed", "64 x 256"],
                   ["Transformer", "6x [PreNorm\u2192MHSA(8 heads)\u2192PreNorm\u2192MLP] + DropPath", "64 x 256"],
                   ["Head", "LayerNorm \u2192 GAP \u2192 Linear(256\u219210)", "10"],
               ],
               caption="Table 10. Hybrid CNN-ViT architecture (5,358,410 trainable parameters).")

    heading(doc, "7.5 Training Results", level=2)
    make_table(doc,
               headers=["Metric", "Value"],
               rows=[
                   ["Trainable parameters", "5,358,410"],
                   ["Training epochs", "120"],
                   ["Best val accuracy", "90.95%"],
                   ["Final train accuracy", "95.8%"],
                   ["Training time", "~99 minutes (5,947s)"],
               ],
               caption="Table 11. Experiment 4 results (real Colab training).")

    add_image(doc, "exp4_curves.png", 5.0, "Figure 5. Experiment 4: Hybrid CNN-ViT training curves.")
    add_image(doc, "exp4_training_dashboard.png", 5.5,
              "Figure 6. Hybrid CNN-ViT comprehensive training dashboard showing loss, accuracy, "
              "LR schedule, generalization gap, cross-experiment comparison, and training efficiency.")

    body(doc,
         "Analysis: The Hybrid CNN-ViT reached 90.95% validation accuracy, a massive improvement "
         "over the pure ViT (69.55%) but slightly below the CNN-only models (92.15% and 92.76%). "
         "The CNN stem successfully addressed the ViT's main weakness by providing pre-extracted "
         "local features. The model shows steady convergence with the generalization gap growing "
         "from ~0% to ~5% over 120 epochs, indicating mild overfitting in later epochs. The "
         "OneCycleLR schedule peaked around epoch 35 and gradually annealed to near-zero by epoch 120.")

    # ═════════════════════════════════════════════════════════════════════
    # 8. RESULTS AND COMPARISON
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "8. Results and Comparison")

    make_table(doc,
               headers=["Model", "Parameters", "Epochs", "Best Val Acc", "Train Time"],
               rows=[
                   ["Baseline CNN", "2,658,762", "100", "92.15%", "68 min"],
                   ["Residual CNN", "6,009,930", "60", "92.76%", "42 min"],
                   ["Vision Transformer", "2,693,194", "40", "69.55%", "31 min"],
                   ["Hybrid CNN-ViT", "5,358,410", "120", "90.95%", "99 min"],
               ],
               caption="Table 12. Final results comparison across all experiments (real Colab training).")

    add_image(doc, "model_comparison.png", 5.0, "Figure 7. Model accuracy and parameter comparison.")
    add_image(doc, "comparison_curves.png", 5.5, "Figure 8. Cross-experiment validation curves showing real training trajectories.")

    heading(doc, "8.1 Key Observations", level=2)
    bullet(doc,
           "The Residual CNN achieved the highest accuracy (92.76%) with the fewest training epochs (60), "
           "demonstrating that well-designed CNNs with skip connections are extremely effective on CIFAR-10.")
    bullet(doc,
           "The Baseline CNN reached 92.15% thanks to the strong training recipe (EMA, label smoothing, "
           "RandAugment), showing that modern training techniques can compensate for simple architectures.")
    bullet(doc,
           "The pure ViT dramatically underperformed at 69.55%, confirming that Transformers need "
           "either large-scale pre-training or architectural modifications to work on small datasets.")
    bullet(doc,
           "The Hybrid CNN-ViT (90.95%) successfully bridged the gap between ViT and CNN, proving "
           "that a CNN stem is essential for making Transformers work on small-scale vision tasks.")

    heading(doc, "8.2 Confusion Matrix Analysis", level=2)
    add_image(doc, "confusion_matrix.png", 4.5, "Figure 9. Confusion matrix for the best model.")
    add_image(doc, "per_class_accuracy.png", 5.0, "Figure 10. Per-class accuracy breakdown.")

    body(doc,
         "The most common misclassifications occur between visually similar classes: "
         "cat/dog (fur texture and body shape overlap), automobile/truck (similar silhouettes), "
         "and bird/airplane (similar outlines at 32\u00d732 resolution). The hardest class is 'cat' "
         "and the easiest classes are 'truck' and 'automobile', which have distinctive geometric shapes.")

    # ═════════════════════════════════════════════════════════════════════
    # 9. DISCUSSION
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "9. Discussion")

    heading(doc, "9.1 Why Residual Connections Help", level=2)
    body(doc,
         "The +0.61% improvement from Exp 1 to Exp 2 demonstrates the benefit of residual "
         "connections, but more importantly, the Residual CNN achieved this in 40% fewer epochs. "
         "Residual connections create shortcut pathways for gradient flow: instead of learning "
         "F(x) from scratch, each block learns the residual R(x) = F(x) - x, which is easier "
         "to optimize. The identity shortcut provides a guaranteed baseline, and the block only "
         "needs to learn refinements. This is why deeper residual networks converge faster and "
         "achieve higher training accuracy (96.6% vs 93.2%).")

    heading(doc, "9.2 Why Pure ViT Failed", level=2)
    body(doc,
         "The pure ViT achieved only 69.55%, a striking 23% below the CNNs. This is the central "
         "finding: Vision Transformers lack the inductive biases of CNNs (translation equivariance "
         "and locality). On CIFAR-10, with only 50K images at 32\u00d732 pixels, the ViT must learn "
         "these patterns entirely from data. With 64 patches of 4\u00d74 pixels each, there is "
         "very limited spatial information per patch. The model was still improving at epoch 40 "
         "(not yet converged), suggesting that with 200-300 epochs and stronger augmentation "
         "(Mixup, CutMix), the ViT could reach ~80-85%, but would still lag behind CNNs.")

    heading(doc, "9.3 The Value of Hybrid Architecture", level=2)
    body(doc,
         "The Hybrid CNN-ViT (90.95%) demonstrates that a CNN stem rescues the Transformer "
         "from its data-efficiency problem. The +21.4% improvement over the pure ViT is dramatic: "
         "the CNN stem provides pre-extracted local features (edges, textures) that the Transformer "
         "can immediately use for global reasoning, instead of spending capacity learning basic "
         "visual patterns. However, the Hybrid still falls slightly short of the pure CNNs (by ~1.8%), "
         "suggesting that on CIFAR-10's scale, the added complexity of Transformer layers does not "
         "yet outweigh the efficiency of well-designed residual CNNs.")

    heading(doc, "9.4 Training Techniques", level=2)
    body(doc,
         "The strong baseline results (92.15% from a simple VGG-style CNN) demonstrate that "
         "modern training techniques provide a large boost:")
    bullet(doc,
           "EMA (Exponential Moving Average): Smoothes weight updates by maintaining a running "
           "average, typically providing +0.3-0.5% improvement.")
    bullet(doc,
           "OneCycleLR: The warmup phase prevents early divergence (especially important for "
           "Transformers), and the cosine annealing phase helps the optimizer find flatter minima.")
    bullet(doc,
           "Label smoothing (0.1): Prevents the model from becoming overconfident on training "
           "examples, improving generalization by ~0.3%.")
    bullet(doc,
           "RandAugment: Provides diverse augmentations that significantly reduce overfitting, "
           "effectively expanding the training set.")
    bullet(doc,
           "Gradient clipping (max_norm=1.0): Essential for Transformer training stability, "
           "prevents gradient spikes from the attention mechanism.")

    heading(doc, "9.5 Overfitting Analysis", level=2)
    body(doc,
         "The generalization gap (train acc - val acc) provides insight into overfitting:")
    make_table(doc,
               headers=["Model", "Final Train Acc", "Best Val Acc", "Gap"],
               rows=[
                   ["Baseline CNN", "93.2%", "92.15%", "1.05%"],
                   ["Residual CNN", "96.6%", "92.76%", "3.84%"],
                   ["ViT", "64.0%", "69.55%", "-5.55% (underfitting)"],
                   ["Hybrid CNN-ViT", "95.8%", "90.95%", "4.85%"],
               ],
               caption="Table 13. Generalization gap analysis.")

    body(doc,
         "The Baseline CNN has the smallest gap (1.05%), indicating excellent regularization from "
         "the heavy Dropout(0.5). The Residual CNN and Hybrid show moderate overfitting (3-5%), "
         "common for larger models on CIFAR-10. The ViT actually shows underfitting (val > train by 5.5%), "
         "meaning it has not yet learned the training distribution adequately in 40 epochs.")

    # ═════════════════════════════════════════════════════════════════════
    # 10. CONCLUSION
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "10. Conclusion")
    body(doc,
         "This study demonstrates the relative strengths and weaknesses of CNN, Transformer, and "
         "Hybrid architectures on the CIFAR-10 image classification task:")

    body(doc, "  Baseline CNN: 92.15% (100 epochs) \u2014 Strong baseline from modern training recipe", bold=True)
    body(doc, "  Residual CNN: 92.76% (60 epochs)  \u2014 Best accuracy, fastest convergence", bold=True)
    body(doc, "  Pure ViT:     69.55% (40 epochs)  \u2014 Insufficient data for attention-only approach", bold=True)
    body(doc, "  Hybrid CNN-ViT: 90.95% (120 epochs) \u2014 CNN stem rescues Transformer performance", bold=True)

    body(doc,
         "Key findings:")
    bullet(doc,
           "Residual connections + modern training recipes make even simple CNNs highly competitive "
           "on CIFAR-10, achieving 92.76% with the Residual CNN.")
    bullet(doc,
           "Pure Vision Transformers fail dramatically (69.55%) on small-scale datasets without "
           "pre-training, confirming their dependence on large data volumes.")
    bullet(doc,
           "Hybrid architectures (90.95%) successfully combine CNN local feature extraction with "
           "Transformer global attention, recovering most of the CNN performance while incorporating "
           "self-attention capabilities.")
    bullet(doc,
           "On CIFAR-10's scale (50K training images, 32\u00d732 pixels), well-designed CNNs still "
           "outperform Transformer-based models, but the hybrid approach shows the path forward "
           "for larger-scale tasks where global attention becomes more valuable.")

    heading(doc, "10.1 Future Work", level=2)
    body(doc,
         "To improve the Hybrid model and close the gap with CNNs, several directions could be explored:")
    bullet(doc, "CutMix/Mixup augmentation: Expected +2-4% improvement by creating more diverse training samples.")
    bullet(doc, "Longer training for ViT and Hybrid: 200-300 epochs with cosine warmup may help Transformers converge.")
    bullet(doc, "Knowledge distillation: Train the Hybrid using a pre-trained CNN teacher (DeiT-style).")
    bullet(doc, "Convolutional position encoding: Replace learnable position embeddings with depthwise convolutions.")
    bullet(doc, "Scale to CIFAR-100 or ImageNet: The Hybrid architecture's advantages may become more pronounced on harder tasks.")

    # ═════════════════════════════════════════════════════════════════════
    # 11. PROMPTS USED TO GENERATE THE NOTEBOOK
    # ═════════════════════════════════════════════════════════════════════
    heading(doc, "11. Prompts Used to Generate the Notebook")

    body(doc,
         "This section documents the prompts and instructions given to AI tools (Claude) "
         "during the development of the Jupyter notebook and codebase. The prompts reflect "
         "the iterative thought process behind the experimental design and implementation.")

    heading(doc, "11.1 Initial Setup & Architecture Design", level=2)

    body(doc, "Prompt 1 — Project Initialization:", bold=True)
    body(doc,
         '"Create a CIFAR-10 image classification study comparing four architectures: '
         '(1) a baseline VGG-style CNN, (2) a CNN with residual connections, '
         '(3) a pure Vision Transformer (ViT), and (4) a Hybrid CNN-ViT model. '
         'Use identical training conditions across all experiments (AdamW optimizer, '
         'OneCycleLR scheduler, EMA, label smoothing, gradient clipping) to ensure '
         'fair comparison. Structure the code in a single Jupyter notebook with clear '
         'sections for each experiment."')

    body(doc, "Prompt 2 — Training Infrastructure:", bold=True)
    body(doc,
         '"Implement a generic experiment runner function that handles: mixed-precision '
         'training (AMP), exponential moving average of weights, OneCycleLR scheduling, '
         'gradient clipping, and label smoothing. The function should track train/val '
         'loss, accuracy, and learning rate per epoch, and save the best model state. '
         'Make it reusable across all four experiments with only the model architecture '
         'changing between runs."')

    heading(doc, "11.2 Architecture Implementation Prompts", level=2)

    body(doc, "Prompt 3 — Baseline CNN:", bold=True)
    body(doc,
         '"Implement Experiment 1: a VGG-style baseline CNN for CIFAR-10. Use three '
         'convolutional blocks with increasing channels (64, 128, 256), each with '
         'BatchNorm + ReLU + MaxPool. Use a fully-connected classifier head with '
         'Dropout(0.5). Train for 100 epochs."')

    body(doc, "Prompt 4 — Residual CNN:", bold=True)
    body(doc,
         '"Implement Experiment 2: an improved CNN with ResNet-style skip connections. '
         'Use a stem conv followed by four residual blocks (128, 256, 256, 512 channels) '
         'with stride-2 downsampling. Replace the FC head with Global Average Pooling '
         'followed by a single linear layer. Add Dropout(0.3). Train for 60 epochs — '
         'residual connections should enable faster convergence."')

    body(doc, "Prompt 5 — Vision Transformer:", bold=True)
    body(doc,
         '"Implement Experiment 3: a pure Vision Transformer for CIFAR-10. Use patch_size=4 '
         '(giving 64 tokens from 32x32 images), embed_dim=192, 6 transformer layers with '
         '3 attention heads and MLP ratio of 4. Include learnable positional embeddings, '
         'stochastic depth (0.1), and dropout (0.1). Use mean pooling over tokens for '
         'classification. Train for 40 epochs."')

    body(doc, "Prompt 6 — Hybrid CNN-ViT:", bold=True)
    body(doc,
         '"Implement Experiment 4: a Hybrid CNN + Vision Transformer. Use a 2-layer CNN stem '
         '(Conv 3→64→128 with BatchNorm + GELU) to extract local features, then divide into '
         '4x4 patches and feed into a 6-layer Transformer encoder (dim=256, 8 heads). The CNN '
         'stem provides local inductive bias that the pure ViT lacks. Train for 120 epochs."')

    heading(doc, "11.3 Visualization & Analysis Prompts", level=2)

    body(doc, "Prompt 7 — Training Curves:", bold=True)
    body(doc,
         '"Generate training visualization plots for each experiment showing: (1) training and '
         'validation loss curves, (2) training and validation accuracy curves, (3) learning rate '
         'schedule. Also create cross-experiment comparison plots overlaying all four models\' '
         'validation accuracy and loss curves. Include a bar chart comparing final accuracies '
         'and parameter counts."')

    body(doc, "Prompt 8 — Confusion Matrix & Per-Class Analysis:", bold=True)
    body(doc,
         '"Generate a confusion matrix heatmap for the best model and a per-class accuracy '
         'bar chart. Analyze which classes are most commonly confused and why (e.g., cat/dog '
         'visual similarity, automobile/truck similar silhouettes)."')

    body(doc, "Prompt 9 — Training Dashboard:", bold=True)
    body(doc,
         '"Create a comprehensive 6-panel training dashboard for Experiment 4 (Hybrid CNN-ViT) '
         'showing: loss curves, accuracy curves, LR schedule, generalization gap over time, '
         'cross-experiment comparison, and a training efficiency metric (accuracy per epoch). '
         'Save as exp4_training_dashboard.png."')

    heading(doc, "11.4 Report & Documentation Prompts", level=2)

    body(doc, "Prompt 10 — Report Generation:", bold=True)
    body(doc,
         '"Generate a comprehensive .docx report with professional formatting including: '
         'title page, introduction with thought process, dataset description, methodology, '
         'detailed sections for each experiment (architecture tables, training results, analysis), '
         'cross-experiment comparison with tables and figures, discussion of why each architecture '
         'performed as it did, and a conclusion with future work directions. Include all generated '
         'plots as embedded figures with captions."')

    body(doc, "Prompt 11 — Iterative Refinement:", bold=True)
    body(doc,
         '"Execute the training pipeline on Google Colab, collect real training logs, and '
         'update the notebook and report with actual results. Regenerate all plots using the '
         'real epoch-by-epoch training data. Verify that the confusion matrix, per-class '
         'accuracy, and all stated metrics match the actual training output."')

    body(doc, "Prompt 12 — Final Quality Check:", bold=True)
    body(doc,
         '"Review and verify the entire submission: check that the notebook has detailed '
         'documentation with markdown cells explaining each section, that all plots are correct '
         'and match the training data, that the report includes tables and written explanations '
         'of the thought process behind each architectural decision, and that the prompts used '
         'to generate the notebook are documented."')

    # ── Save ────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Report saved \u2192 {OUT_PATH}")


if __name__ == "__main__":
    build_report()
